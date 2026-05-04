# -*- coding: utf-8 -*-
"""
gen_docx_v1.py — AI컨설팅시스템_종합역량문서_v1.docx 생성
실행: python outputs/gen_docx_v1.py
"""
import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = pathlib.Path(r"C:\Users\Jy\consulting-agent\outputs")

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="000080", sz=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    default_sz = {1: 16, 2: 13, 3: 11}
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
        run.font.size = Pt(sz or default_sz.get(level, 11))
        run.font.name = "Malgun Gothic"

def add_para(doc, text, sz=9.5, bold=False, color="1E293B", indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.name = "Malgun Gothic"
    if color != "1E293B":
        r.font.color.rgb = RGBColor.from_string(color)
    return p

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        r = cell.paragraphs[0].runs[0]
        r.font.bold = True; r.font.size = Pt(8.5); r.font.name = "Malgun Gothic"
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        set_cell_bg(cell, "000080")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = tbl.add_row()
        for i, v in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(v)
            r = cell.paragraphs[0].runs[0]
            r.font.size = Pt(8.5); r.font.name = "Malgun Gothic"
    return tbl

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)

# 표제부
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("AI 컨설팅 시스템 종합역량 문서")
r.font.size = Pt(22); r.font.bold = True
r.font.color.rgb = RGBColor.from_string("000080"); r.font.name = "Malgun Gothic"

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("중기이코노미기업지원단 | 전문위원 여운식 | 2026-05-04 | v1.0")
r2.font.size = Pt(9.5); r2.font.name = "Malgun Gothic"
r2.font.color.rgb = RGBColor.from_string("64748B")
doc.add_paragraph()

# ──────────────────────────────────────────────
# 챕터 1: 코워크 모드 스킬
# ──────────────────────────────────────────────
add_heading(doc, "챕터 1: 코워크 모드 스킬 (Claude Code Skills)")
add_para(doc, "Claude Code 스킬 시스템은 .skills/ 디렉토리에 SKILL.md 파일로 정의된다. 각 스킬은 독립 트리거·역할·출력 형식을 가지며, 코워크 모드에서 자동 라우팅된다.")

add_heading(doc, "1.1 등록 스킬 목록 (5종)", level=2)
add_table(doc,
    ["스킬명", "수준★", "역할", "트리거 키워드"],
    [
        ["corp-financial-analysis", "★★★★★", "6트랙 재무·세무 분석", "법인세 절세, 재무 분석"],
        ["pptx", "★★★★", "네이비 표준 PPT 생성", "PPT 만들어, 슬라이드"],
        ["docx", "★★★★", "표준 Word 보고서 생성", "Word 보고서, docx"],
        ["xlsx", "★★★★", "시뮬레이터·분석 Excel", "Excel 만들어, 시뮬"],
        ["pdf", "★★★", "최종 배포용 PDF 변환", "PDF 변환, 인쇄용"],
    ]
)
doc.add_paragraph()

# ──────────────────────────────────────────────
# 챕터 2: MCP 플러그인
# ──────────────────────────────────────────────
add_heading(doc, "챕터 2: MCP 플러그인")
add_para(doc, "RULE-G2 준수: 쓰기성 MCP 도구는 사용자 승인 후에만 호출한다.")

add_heading(doc, "2.1 읽기성 도구 (자동 허용)", level=2)
add_table(doc,
    ["도구명", "기능"],
    [
        ["Gmail search_threads / get_thread", "메일 검색 및 조회"],
        ["Google Drive search / read / list", "Drive 파일 탐색·읽기"],
        ["Google Calendar authenticate", "캘린더 인증"],
    ]
)
doc.add_paragraph()

add_heading(doc, "2.2 쓰기성 도구 (RULE-G2 — 승인 필수)", level=2)
add_table(doc,
    ["도구명", "기능", "승인 사유"],
    [
        ["Gmail create_draft", "메일 초안 생성", "발송 부주의 방지"],
        ["Drive create_file / copy_file", "파일 생성·복사", "의도치 않은 생성 방지"],
        ["Calendar complete_authentication", "인증 완료", "권한 취득 주의"],
    ]
)
doc.add_paragraph()

# ──────────────────────────────────────────────
# 챕터 3: 하네스 에이전트팀
# ──────────────────────────────────────────────
add_heading(doc, "챕터 3: 하네스 에이전트팀")

add_heading(doc, "3.1 에이전트 현황 (2026-05-04 기준)", level=2)
add_table(doc,
    ["구분", "수량", "위치"],
    [
        ["Active 에이전트", "33개", "agents/active/*.py"],
        ["Proposals (명세)", "40개", "agents/proposals/*.md"],
        ["합계", "73개", "—"],
    ]
)
doc.add_paragraph()

add_heading(doc, "3.2 그룹별 분류", level=2)
add_table(doc,
    ["그룹", "대표 에이전트", "역할"],
    [
        ["분석·진단 (8개)", "financial_statement_pdf_parser, kredtop_parser", "재무·신용 데이터 파싱"],
        ["절세·승계 (8개)", "inheritance_gift_agent, child_corp_design", "상속·증여·승계 최적화"],
        ["PART6 신규 (4개)", "business_plan_agent, legal_risk_hedge", "사업계획·법무·R&D"],
        ["시스템·자가진화 (5개)", "discovery_agent, verifier_agent", "자가 발견·실행·검증"],
        ["기타 (8개)", "ppt_polisher, pii_masking_agent", "산출물·보안 처리"],
    ]
)
doc.add_paragraph()

add_heading(doc, "3.3 목표: 93+ 에이전트 달성 로드맵", level=2)
add_table(doc,
    ["단계", "추가 대상", "예상 수량"],
    [
        ["D+7", "정책자금·보증 2종", "+2"],
        ["D+14", "DART 공시 분석 3종", "+3"],
        ["D+21", "법령 자동 모니터링 3종", "+3"],
        ["D+30", "인사노무 전용 4종", "+4"],
        ["D+45", "기타 전문 분야 6종", "+6"],
        ["합계", "73 + 18", "= 91개"],
    ]
)
doc.add_paragraph()

# ──────────────────────────────────────────────
# 챕터 4: 업무역량 향상 추천
# ──────────────────────────────────────────────
add_heading(doc, "챕터 4: 업무역량 향상 추천")

add_heading(doc, "4.1 즉시 실행 (무료 범위 6건)", level=2)
add_table(doc,
    ["#", "항목", "방법", "기대 효과"],
    [
        ["1", "/마스터컨설팅 슬래시명령", ".claude/commands/ 등록 완료", "종합 분석 1회 실행"],
        ["2", "VS Code Ctrl+Shift+B", ".vscode/tasks.json 10개", "반복 작업 자동화"],
        ["3", "6트랙 병렬 분석", "corp-financial-analysis 스킬", "분석 시간 60% 단축"],
        ["4", "4축 검증 매트릭스", "산출물 생성 후 자동 적용", "오류 0건 달성"],
        ["5", "sync_all_devices.ps1", "3대 기기 1분 동기화", "기기 간 일관성 유지"],
        ["6", "RULE-G 자가검증", "세션 종료 전 3문항", "안전사고 재발 방지"],
    ]
)
doc.add_paragraph()

add_heading(doc, "4.2 신규 보강 추천 (D+7 이내 3건)", level=2)
add_table(doc,
    ["#", "항목", "내용", "예상시간"],
    [
        ["1", "/진단요약 신설", "4자관점 1페이지 요약 자동 생성", "15분"],
        ["2", "법인재무분석 SKILL.md", "코워크모드 직접 호출 스킬", "20분"],
        ["3", "pre_session_check.ps1", "세션 시작 전 환경 자동 점검", "30분"],
    ]
)
doc.add_paragraph()

# ──────────────────────────────────────────────
# 챕터 5: CLAUDE.md 보강
# ──────────────────────────────────────────────
add_heading(doc, "챕터 5: CLAUDE.md 보강 전후")

add_heading(doc, "5.1 보강 전후 비교", level=2)
add_table(doc,
    ["항목", "보강 전", "보강 후"],
    [
        ["git push 안전장치", "없음", "RULE-G1 — 사용자 명시 없는 push 금지"],
        ["MCP 쓰기성 도구", "자동 호출", "RULE-G2 — 호출 전 사용자 승인 필수"],
        ["idle timeout 방지", "없음", "RULE-G3 — 5파일/2,000토큰 배치 분할"],
        ["자가검증 기준", "산출물별 부분", "세션 종료 전 3축 자가질문"],
        ["슬래시명령 위치", "junggi-workspace만", "consulting-agent .claude/commands/"],
        ["스킬 파일", "없음", ".skills/ 5종 체계화"],
    ]
)
doc.add_paragraph()

add_heading(doc, "5.2 정량 지표 (보강 후)", level=2)
add_table(doc,
    ["지표", "값"],
    [
        ["RULE-G 섹션", "3조 (G1·G2·G3)"],
        ["등록 슬래시명령", "10종 (.claude/commands/)"],
        ["등록 스킬", "5종 (.skills/)"],
        ["VS Code task", "10개"],
        ["동기화 스크립트", "2종 (ps1·sh)"],
        ["자가검증 축", "3축 (push·MCP·배치)"],
    ]
)
doc.add_paragraph()

# 엔딩
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = end_p.add_run("감사합니다")
r_end.font.size = Pt(22); r_end.font.bold = True
r_end.font.color.rgb = RGBColor.from_string("000080"); r_end.font.name = "Malgun Gothic"

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_c = contact_p.add_run("영남사업단 경남1본부(진주센터) 전문위원 여운식 | 010-2673-3781 | sik3781@gmail.com")
r_c.font.size = Pt(9); r_c.font.name = "Malgun Gothic"
r_c.font.color.rgb = RGBColor.from_string("64748B")

out_path = OUT / "AI컨설팅시스템_종합역량문서_v1.docx"
doc.save(out_path)
print(f"✅ {out_path} ({out_path.stat().st_size:,} bytes)")
