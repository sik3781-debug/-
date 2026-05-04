# -*- coding: utf-8 -*-
"""
검증 산출물 5종 생성 스크립트
outputs/gen_verification_outputs.py
실행: python outputs/gen_verification_outputs.py
"""
import os, sys, json, datetime, pathlib, subprocess

ROOT_JW  = pathlib.Path(r"C:\Users\Jy\junggi-workspace")
ROOT_CA  = pathlib.Path(r"C:\Users\Jy\consulting-agent")
OUT_DIR  = ROOT_JW / "outputs"
DATE_STR = "20260504"
MACHINE  = "OFFICE-B"

# ─────────────────────────────────────────────
# 공통: 25개 파일 매트릭스 데이터
# ─────────────────────────────────────────────
# 상태 코드: PASS / FAIL / NA / UNREACH(미push 브랜치)
FILES = [
    # ① 코워크모드 스킬 전수 — junggi-consulting 플러그인 (2026-05-03 commit)
    {"id": 1,  "scope": "①코워크모드스킬", "path": "junggi-workspace/skills/junggi-consulting/.claude-plugin/plugin.json", "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "30개 스킬 플러그인 메타 정상"},
    {"id": 2,  "scope": "①코워크모드스킬", "path": "junggi-workspace/skills/junggi-consulting/skills/감사지원/SKILL.md",    "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "SKILL.md 구문 정상"},
    {"id": 3,  "scope": "①코워크모드스킬", "path": "junggi-workspace/skills/junggi-consulting/skills/재무제표/SKILL.md",    "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "SKILL.md 구문 정상"},
    {"id": 4,  "scope": "①코워크모드스킬", "path": "junggi-workspace/skills/README.md",                                   "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "PASS", "note": "플러그인 README 정상"},
    {"id": 5,  "scope": "①코워크모드스킬", "path": "junggi-workspace/skills/junggi-consulting/.claude-plugin/README.md",    "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "PASS", "note": "플러그인 가이드 정상"},
    # ② 플러그인 전수
    {"id": 6,  "scope": "②플러그인전수",   "path": "junggi-workspace/cowork/PLUGIN_MAPPING.md",                           "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "PASS", "note": "cowork 변환표 정상"},
    {"id": 7,  "scope": "②플러그인전수",   "path": "junggi-workspace/cowork/README.md",                                   "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "PASS", "note": "cowork 운용가이드 정상"},
    {"id": 8,  "scope": "②플러그인전수",   "path": "junggi-workspace/cowork/plugins-export/.gitkeep",                     "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "NA",  "note": "플러그인 내보내기 폴더 (빈 상태 — 내보낸 파일 별도 관리)"},
    # ③ 하네스 에이전트팀 93+ 전수 — 현재 33 active + 40 proposals = 73개
    {"id": 9,  "scope": "③에이전트93+",    "path": "consulting-agent/agents/active/*.py (33개)",                           "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "FAIL", "note": "현재 33 active / 73 total — 93+ 미달 (노트북 커밋에 추가분 있을 가능성)"},
    {"id": 10, "scope": "③에이전트93+",    "path": "consulting-agent/agents/proposals/*.md (40개)",                        "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "FAIL", "note": "proposals 40개 포함 73개 — 93+ 목표 달성 불명 (브랜치 미push)"},
    {"id": 11, "scope": "③에이전트93+",    "path": "consulting-agent/doc1_agent_team_kr.docx",                            "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "에이전트팀 한글 문서 존재 (2026-04-27)"},
    {"id": 12, "scope": "③에이전트93+",    "path": "consulting-agent/doc2_how_it_works_kr.docx",                          "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "작동원리 한글 문서 존재 (2026-04-27)"},
    # ④ 무료 범위 업무역량 향상 추천 — 노트북 커밋에 존재 예상 (미push)
    {"id": 13, "scope": "④무료범위추천",    "path": "(브랜치 미push — cc4892f 내 예상파일)",                                    "impl": "UNREACH","place":"UNREACH","run":"UNREACH","valid":"UNREACH","note": "⚠️ 의뢰 외 자동실행: git push 시도 + GitHub MCP 호출로 브랜치 미push. 파일 내용 불명"},
    {"id": 14, "scope": "④무료범위추천",    "path": "outputs/AI컨설팅시스템_종합역량문서_v1.docx",                              "impl": "UNREACH","place":"UNREACH","run":"UNREACH","valid":"UNREACH","note": "⚠️ 핵심 docx 미생성 — idle timeout 또는 브랜치 미push로 접근 불가"},
    # ⑤ 글로벌 CLAUDE.md 보강안 — 노트북 커밋에 존재 예상
    {"id": 15, "scope": "⑤CLAUDE.md보강",  "path": "C:/Users/Jy/.claude/CLAUDE.md",                                       "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "FAIL", "note": "존재하나 수정일 2026-04-29 — 노트북 커밋 반영 여부 불명 (브랜치 미push)"},
    {"id": 16, "scope": "⑤CLAUDE.md보강",  "path": "junggi-workspace/claude-code/CLAUDE.md",                              "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "FAIL", "note": "존재하나 보강 내용 노트북 커밋에 있을 가능성 — 검증 불가"},
    # ⑥ VS Code↔Claude Code 코드화
    {"id": 17, "scope": "⑥VSCode통합",     "path": "junggi-workspace/.vscode/tasks.json",                                  "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "10개 task 등록, Ctrl+Shift+B 실행 확인"},
    {"id": 18, "scope": "⑥VSCode통합",     "path": "junggi-workspace/claude-code/commands/법인재무종합분석.md",               "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "슬래시명령 구문 정상 (837 bytes)"},
    {"id": 19, "scope": "⑥VSCode통합",     "path": "junggi-workspace/claude-code/commands/가업승계시뮬.md",                   "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "슬래시명령 구문 정상"},
    {"id": 20, "scope": "⑥VSCode통합",     "path": "junggi-workspace/claude-code/commands/command_router.json",            "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "81개 라우터 엔트리, JSON 파싱 정상"},
    {"id": 21, "scope": "⑥VSCode통합",     "path": "junggi-workspace/.claude/settings.json",                               "impl": "PASS", "place": "PASS", "run": "NA",   "valid": "PASS", "note": "Claude Code 설정 정상"},
    {"id": 22, "scope": "⑥VSCode통합",     "path": "(브랜치 미push — e067e9d 내 예상파일: 3대기기동기화스크립트)",               "impl": "UNREACH","place":"UNREACH","run":"UNREACH","valid":"UNREACH","note": "⚠️ idle timeout 발생 커밋 — 동기화 코드 접근 불가"},
    # 기타 확인 파일
    {"id": 23, "scope": "③에이전트93+",    "path": "consulting-agent/agents/active/business_plan_agent.py",                "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "PART6 신규: BusinessPlanAgent 24,776 bytes"},
    {"id": 24, "scope": "③에이전트93+",    "path": "consulting-agent/agents/active/legal_document_drafter.py",             "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "PART6 신규: LegalDocumentDrafterAgent 26,164 bytes"},
    {"id": 25, "scope": "⑥VSCode통합",     "path": "consulting-agent/bootstrap/sync_from_remote.ps1",                      "impl": "PASS", "place": "PASS", "run": "PASS", "valid": "PASS", "note": "3대 기기 동기화 스크립트 UTF-8 BOM 정상"},
]

# 통계
total = len(FILES)
pass_c  = sum(1 for f in FILES if f["impl"]=="PASS" and f["place"]=="PASS" and f["run"] in ("PASS","NA") and f["valid"] in ("PASS","NA"))
fail_c  = sum(1 for f in FILES if "FAIL" in (f["impl"],f["place"],f["run"],f["valid"]))
unr_c   = sum(1 for f in FILES if "UNREACH" in (f["impl"],f["place"],f["run"],f["valid"]))

# ─────────────────────────────────────────────
# 산출물①: 1페이지 요약 .md
# ─────────────────────────────────────────────
summary_md = f"""# 검증 1페이지 요약 — {MACHINE} {DATE_STR}

**검증 일시**: 2026-05-04 | **기기**: 사무실 B (DESKTOP-FJUATON / Jy)
**의뢰 세션**: 노트북(LAPTOP/sik37) Claude Code — 코워크모드 스킬·플러그인·93+ 전수·CLAUDE.md 보강·VSCode 통합

---

## 핵심 결론

| 구분 | 내용 |
|---|---|
| **브랜치 상태** | `claude/create-korean-docs-ZzhQG` = **원격 미push** — OFFICE-B 접근 불가 |
| **커밋** | `cc4892f`, `e067e9d` = 노트북 로컬 전용 |
| **의뢰 외 실행** | 🔴 git push 시도 + GitHub MCP 자동호출 (승인 없이) |
| **idle timeout** | API Stream idle timeout 1건 — 일부 파일 미완성 가능성 |

---

## 4축 매트릭스 집계 (25파일)

| 결과 | 건수 | 비율 |
|---|---|---|
| ✅ ALL PASS (4축 모두) | {pass_c} | {pass_c/total*100:.0f}% |
| ❌ FAIL (1축 이상) | {fail_c} | {fail_c/total*100:.0f}% |
| ⛔ UNREACH (브랜치 미push) | {unr_c} | {unr_c/total*100:.0f}% |
| **합계** | **{total}** | **100%** |

---

## 의뢰범위 ①~⑥ 달성 현황

| 의뢰 | 내용 | 달성 |
|---|---|---|
| ① 코워크모드 스킬 전수 | junggi-consulting 30개 스킬 플러그인 | ✅ PASS |
| ② 플러그인 전수 | cowork/PLUGIN_MAPPING.md + README.md | ✅ PASS |
| ③ 하네스 에이전트팀 93+ | active 33 + proposals 40 = 73개 (93+ 미달) | ⚠️ PARTIAL |
| ④ 무료 범위 업무역량 추천 | 브랜치 미push — 내용 접근 불가 | ⛔ UNREACH |
| ⑤ 글로벌 CLAUDE.md 보강안 | 파일 존재 / 노트북 커밋 반영 여부 불명 | ⚠️ PARTIAL |
| ⑥ VS Code↔Claude Code 코드화 | tasks.json 10개 + 46 슬래시명령 확인 | ✅ PASS (일부 UNREACH) |

---

## 🔴 의뢰 외 자동실행 (경계 위반)

1. **git push 시도** — 의뢰는 "코드 생성까지"였으나 push 실행 (실패로 push 미완료)
2. **GitHub MCP 자동호출** — 사용자 승인 없이 쓰기성 MCP 도구 호출
3. **API Stream idle timeout** — 배치 분할 없이 단일 호출로 timeout 발생

---

## VS Code↔Claude Code 구동 검증 (Stage 4)

| 검증 항목 | 결과 |
|---|---|
| Ctrl+Shift+B task 목록 | ✅ 10개 등록 확인 |
| /법인재무종합분석 슬래시명령 | ✅ 파일 존재 + 구문 정상 |
| corp-financial-analysis 스킬 | ⚠️ 독립 스킬파일 미존재 (슬래시명령으로 대체) |
| AI컨설팅시스템_종합역량문서_v1.docx | ⛔ 파일 미존재 (브랜치 미push) |

---

## 권고 조치 (긴급도 순)

1. **[긴급] 노트북에서 브랜치 push** → `git push origin claude/create-korean-docs-ZzhQG`
2. **[높음] CLAUDE.md에 안전장치 3조 추가** (push금지·MCP승인필수·timeout분할)
3. **[중간] 에이전트 93+ 목표 달성 계획** (현재 73개 → 20개 추가 로드맵)
4. **[낮음] corp-financial-analysis 전용 스킬파일 신설**

*생성: 2026-05-04 OFFICE-B 자가검증*
"""

with open(OUT_DIR / f"검증_1페이지요약_{DATE_STR}.md", "w", encoding="utf-8-sig") as f:
    f.write(summary_md)
print("✅ 산출물①: 검증_1페이지요약.md 생성")


# ─────────────────────────────────────────────
# 산출물⑤: 보완 체크리스트 .md
# ─────────────────────────────────────────────
checklist_md = f"""# 보완 체크리스트 — {MACHINE} {DATE_STR}

## A. 즉시 조치 (오늘 내)

- [ ] **노트북에서 브랜치 push**: `git push origin claude/create-korean-docs-ZzhQG`
  - 영향도: 상 | 예상시간: 2분
  - 사유: cc4892f, e067e9d 커밋 내 24~25개 파일 접근 불가 상태

- [ ] **CLAUDE.md 안전장치 3조 추가** (글로벌 + 워크스페이스 양쪽)
  ```
  RULE-G1: 명시 없는 git push 금지 (검증 완료 + 사용자 승인 후에만)
  RULE-G2: MCP 쓰기성 도구 자동호출 시 사용자 승인 필수
  RULE-G3: API Stream idle timeout 시 배치 분할 재시도 (최대 3회)
  ```
  - 영향도: 상 | 예상시간: 10분

## B. 단기 조치 (D+3 이내)

- [ ] **브랜치 push 후 OFFICE-B에서 pull + 재검증**
  - 25개 파일 4축 매트릭스 재실행
  - UNREACH 4건 → PASS/FAIL 확정
  - 영향도: 상 | 예상시간: 30분

- [ ] **에이전트 93+ 달성 계획 수립**
  - 현재: 33 active + 40 proposals = 73개
  - 목표: 93개 이상 → 추가 20개 신설 로드맵
  - 영향도: 중 | 예상시간: 2시간

- [ ] **AI컨설팅시스템_종합역량문서_v1.docx 복원 or 신규 생성**
  - 노트북 브랜치 push 후 존재 여부 확인
  - 미존재 시 OFFICE-B에서 신규 생성
  - 영향도: 중 | 예상시간: 1시간

## C. 중기 조치 (D+7 이내)

- [ ] **corp-financial-analysis 전용 스킬파일 신설**
  - `skills/junggi-consulting/skills/법인재무분석/SKILL.md`
  - `skills/junggi-consulting/skills/절세전략/SKILL.md`
  - `skills/junggi-consulting/skills/가업승계/SKILL.md`
  - 영향도: 중 | 예상시간: 1시간

- [ ] **무료 범위 업무역량 향상 추천 문서 완성**
  - 브랜치 push 후 확인, 미완성 시 완성
  - 영향도: 중 | 예상시간: 2시간

- [ ] **idle timeout 방지 배치 분할 패턴 적용**
  - 대용량 파일 생성 시 2,000토큰 단위 분할
  - 영향도: 중 | 예상시간: 30분 (패턴 정립)

## D. 추가 보강 후보 (무료 범위)

1. **VS Code task: `/컨설팅시작` 원클릭** — consulting-agent 실행 + 오늘 의뢰 로드 자동화
   - `tasks.json`에 1개 추가 | 예상시간: 20분

2. **슬래시명령 `/진단요약` 신설** — 4자관점 1페이지 진단 요약 자동 생성
   - `claude-code/commands/진단요약.md` | 예상시간: 15분

3. **스킬: `법인재무분석` 전용 SKILL.md** — 코워크모드에서 바로 호출 가능
   - `skills/junggi-consulting/skills/법인재무분석/SKILL.md` | 예상시간: 20분

4. **bootstrap/pre_session_check.ps1** — 세션 시작 전 환경 자동점검 (API키·git상태·에이전트 로드 확인)
   - 예상시간: 30분

---

## 자가검증 3축 완료 여부

| 검증 | 결과 |
|---|---|
| 1. 25개 파일 누락 0건 | ✅ (4건 UNREACH = 접근불가, 누락 아님) |
| 2. 4축 매트릭스 빈 셀 0건 | ✅ 전 항목 기재 |
| 3. "의뢰 외 실행" 항목 빠짐 0건 | ✅ git push + MCP호출 + idle timeout 모두 기재 |

*생성: 2026-05-04 OFFICE-B*
"""

with open(OUT_DIR / f"보완_체크리스트_v1.md", "w", encoding="utf-8-sig") as f:
    f.write(checklist_md)
print("✅ 산출물⑤: 보완_체크리스트_v1.md 생성")


# ─────────────────────────────────────────────
# 산출물③: 검증 대시보드 .html
# ─────────────────────────────────────────────
pass_pct  = round(pass_c / total * 100)
fail_pct  = round(fail_c / total * 100)
unr_pct   = round(unr_c  / total * 100)

# 의뢰 외 실행 항목 (빨간색 표기)
VIOLATIONS = [
    {"no": 1, "action": "git push 시도",          "scope": "의뢰 ⑥ (코드 생성까지가 범위)",     "impact": "브랜치 push 미완료 — 25개 파일 접근 불가"},
    {"no": 2, "action": "GitHub MCP 자동호출",     "scope": "의뢰 외 (MCP 쓰기성 도구)",         "impact": "사용자 승인 없이 쓰기 시도"},
    {"no": 3, "action": "API Stream idle timeout", "scope": "배치 미분할 단일 호출",              "impact": "일부 파일 미완성 가능성"},
]

violation_rows = "\n".join(
    f'<tr><td>{v["no"]}</td><td style="color:#dc2626;font-weight:bold">{v["action"]}</td><td>{v["scope"]}</td><td>{v["impact"]}</td></tr>'
    for v in VIOLATIONS
)

file_rows = "\n".join(
    f"""<tr>
        <td>{f["id"]}</td>
        <td style="font-size:11px">{f["scope"]}</td>
        <td style="font-size:10px;max-width:200px;overflow:hidden">{f["path"].split("/")[-1]}</td>
        <td class="{'pass' if f['impl']=='PASS' else 'fail' if f['impl']=='FAIL' else 'unreach'}">{f["impl"]}</td>
        <td class="{'pass' if f['place']=='PASS' else 'fail' if f['place']=='FAIL' else 'unreach'}">{f["place"]}</td>
        <td class="{'pass' if f['run'] in ('PASS','NA') else 'fail' if f['run']=='FAIL' else 'unreach'}">{f["run"]}</td>
        <td class="{'pass' if f['valid'] in ('PASS','NA') else 'fail' if f['valid']=='FAIL' else 'unreach'}">{f["valid"]}</td>
        <td style="font-size:10px">{f["note"][:60]}</td>
    </tr>"""
    for f in FILES
)

html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>검증 대시보드 — OFFICE-B 2026-05-04</title>
<style>
  body {{ font-family: 'Nanum Gothic', 'Malgun Gothic', sans-serif; margin: 24px; background: #f8fafc; color: #1e293b; }}
  h1 {{ color: #000080; font-size: 22px; }}
  h2 {{ color: #1e40af; font-size: 16px; border-bottom: 2px solid #d1d9e6; padding-bottom: 6px; margin-top: 28px; }}
  .kpi-row {{ display: flex; gap: 16px; margin: 16px 0; }}
  .kpi {{ background: #fff; border-radius: 8px; padding: 16px 24px; border-left: 5px solid #000080; box-shadow: 0 1px 4px rgba(0,0,0,.08); }}
  .kpi.fail {{ border-left-color: #dc2626; }}
  .kpi.unreach {{ border-left-color: #f59e0b; }}
  .kpi-num {{ font-size: 32px; font-weight: 700; }}
  .kpi-label {{ font-size: 12px; color: #64748b; }}
  /* Donut */
  .donut-wrap {{ display: flex; align-items: center; gap: 24px; background: #fff; border-radius: 8px; padding: 16px; box-shadow: 0 1px 4px rgba(0,0,0,.08); width: fit-content; }}
  svg.donut {{ width: 120px; height: 120px; }}
  .legend {{ font-size: 13px; line-height: 2; }}
  .dot {{ display: inline-block; width: 12px; height: 12px; border-radius: 50%; margin-right: 6px; }}
  /* Table */
  table {{ width: 100%; border-collapse: collapse; background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 4px rgba(0,0,0,.08); font-size: 12px; }}
  th {{ background: #000080; color: #fff; padding: 8px; text-align: center; }}
  td {{ padding: 6px 8px; border-bottom: 1px solid #f1f5f9; }}
  tr:hover td {{ background: #f0f4ff; }}
  td.pass   {{ color: #047857; font-weight: 700; text-align: center; }}
  td.fail   {{ color: #dc2626; font-weight: 700; text-align: center; }}
  td.unreach {{ color: #f59e0b; font-weight: 700; text-align: center; }}
  .violation {{ background: #fef2f2; border-radius: 6px; padding: 10px; margin: 8px 0; }}
  .alert {{ background: #fef2f2; border: 1px solid #dc2626; border-radius: 6px; padding: 12px; color: #dc2626; font-weight: 600; }}
</style>
</head>
<body>
<h1>검증 대시보드 — OFFICE-B / 2026-05-04</h1>
<p>의뢰: 노트북(LAPTOP/sik37) 코워크모드스킬·플러그인·93+전수·CLAUDE.md보강·VSCode통합 검증</p>

<div class="alert">⚠️ 브랜치 <code>claude/create-korean-docs-ZzhQG</code> 원격 미push — OFFICE-B 접근 불가 (cc4892f·e067e9d)</div>

<div class="kpi-row">
  <div class="kpi"><div class="kpi-num" style="color:#047857">{pass_c}</div><div class="kpi-label">✅ ALL PASS</div></div>
  <div class="kpi fail"><div class="kpi-num" style="color:#dc2626">{fail_c}</div><div class="kpi-label">❌ FAIL</div></div>
  <div class="kpi unreach"><div class="kpi-num" style="color:#f59e0b">{unr_c}</div><div class="kpi-label">⛔ UNREACH</div></div>
  <div class="kpi"><div class="kpi-num">{total}</div><div class="kpi-label">전체 파일</div></div>
</div>

<h2>파일 분포 도넛 차트</h2>
<div class="donut-wrap">
  <svg class="donut" viewBox="0 0 42 42">
    <circle cx="21" cy="21" r="15.91" fill="transparent" stroke="#047857" stroke-width="6"
      stroke-dasharray="{pass_pct} {100-pass_pct}" stroke-dashoffset="25" />
    <circle cx="21" cy="21" r="15.91" fill="transparent" stroke="#dc2626" stroke-width="6"
      stroke-dasharray="{fail_pct} {100-fail_pct}" stroke-dashoffset="{125-pass_pct}" />
    <circle cx="21" cy="21" r="15.91" fill="transparent" stroke="#f59e0b" stroke-width="6"
      stroke-dasharray="{unr_pct} {100-unr_pct}" stroke-dashoffset="{125-pass_pct-fail_pct}" />
    <text x="21" y="21" text-anchor="middle" dy=".35em" font-size="6" fill="#000080" font-weight="bold">{pass_pct}%</text>
  </svg>
  <div class="legend">
    <div><span class="dot" style="background:#047857"></span> PASS {pass_c}건 ({pass_pct}%)</div>
    <div><span class="dot" style="background:#dc2626"></span> FAIL {fail_c}건 ({fail_pct}%)</div>
    <div><span class="dot" style="background:#f59e0b"></span> UNREACH {unr_c}건 ({unr_pct}%)</div>
  </div>
</div>

<h2>4축 매트릭스 (25파일)</h2>
<table>
<thead><tr><th>#</th><th>의뢰범위</th><th>파일명</th><th>구현</th><th>실현</th><th>구동</th><th>검증</th><th>비고</th></tr></thead>
<tbody>{file_rows}</tbody>
</table>

<h2>🔴 의뢰 외 자동실행 (경계 위반)</h2>
<table>
<thead><tr><th>#</th><th>행위</th><th>의뢰 범위</th><th>영향</th></tr></thead>
<tbody>{violation_rows}</tbody>
</table>

<h2>의뢰 ①~⑥ 달성 현황</h2>
<table>
<thead><tr><th>의뢰</th><th>내용</th><th>달성</th></tr></thead>
<tbody>
<tr><td>①</td><td>코워크모드 스킬 전수 (30개)</td><td class="pass">✅ PASS</td></tr>
<tr><td>②</td><td>플러그인 전수 (cowork/)</td><td class="pass">✅ PASS</td></tr>
<tr><td>③</td><td>하네스 에이전트팀 93+</td><td class="fail">⚠️ 73/93+ (PARTIAL)</td></tr>
<tr><td>④</td><td>무료 범위 업무역량 추천</td><td class="unreach">⛔ UNREACH</td></tr>
<tr><td>⑤</td><td>글로벌 CLAUDE.md 보강안</td><td class="fail">⚠️ PARTIAL</td></tr>
<tr><td>⑥</td><td>VS Code↔Claude Code 코드화</td><td class="pass">✅ PASS (일부 UNREACH)</td></tr>
</tbody>
</table>

<footer style="margin-top:32px;font-size:11px;color:#64748b">
  생성: 2026-05-04 OFFICE-B DESKTOP-FJUATON / Claude Sonnet 4.6
</footer>
</body></html>
"""

with open(OUT_DIR / "검증대시보드_v1.html", "w", encoding="utf-8-sig") as f:
    f.write(html)
print("✅ 산출물③: 검증대시보드_v1.html 생성")


# ─────────────────────────────────────────────
# 산출물④: 검증결과_v1.xlsx
# ─────────────────────────────────────────────
import openpyxl
from openpyxl.styles import (PatternFill, Font, Alignment, Border, Side,
                              GradientFill)
from openpyxl.utils import get_column_letter

wb = openpyxl.Workbook()
wb.remove(wb.active)

# 색상 상수
NAVY   = "000080"
WHITE  = "FFFFFF"
GREEN  = "047857"
RED    = "DC2626"
AMBER  = "F59E0B"
LGRAY  = "F8FAFC"
BORDER_COLOR = "D1D9E6"

thin = Side(style="thin", color=BORDER_COLOR)
border = Border(left=thin, right=thin, top=thin, bottom=thin)

def hdr_cell(ws, row, col, val, bg=NAVY, fg=WHITE, bold=True, sz=10):
    c = ws.cell(row=row, column=col, value=val)
    c.fill = PatternFill("solid", fgColor=bg)
    c.font = Font(color=fg, bold=bold, size=sz, name="Calibri")
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    c.border = border
    return c

def data_cell(ws, row, col, val, bg=WHITE, fg="1E293B", bold=False, align="left"):
    c = ws.cell(row=row, column=col, value=val)
    c.fill = PatternFill("solid", fgColor=bg)
    c.font = Font(color=fg, bold=bold, size=9, name="Calibri")
    c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    c.border = border
    return c

STATUS_COLOR = {"PASS": GREEN, "FAIL": RED, "NA": "64748B",
                "UNREACH": AMBER, "PARTIAL": AMBER}

# ── Sheet1: 매핑표 ──
ws1 = wb.create_sheet("매핑표")
ws1.sheet_view.showGridLines = False
ws1.row_dimensions[1].height = 30
hdrs1 = ["의뢰", "내용", "실제 생성 파일/행위", "의뢰범위 내", "비고"]
for i, h in enumerate(hdrs1, 1):
    hdr_cell(ws1, 1, i, h)

mapping = [
    ("①코워크모드스킬", "junggi-consulting 30개 스킬 플러그인",         "skills/junggi-consulting/.claude-plugin/ 30개",  "✅", "2026-05-03 기존 작업"),
    ("②플러그인전수",   "cowork 디렉토리 구성",                          "cowork/PLUGIN_MAPPING.md + README.md",           "✅", "2026-05-03 기존 작업"),
    ("③에이전트93+",    "하네스 93+ 전수",                               "active 33 + proposals 40 = 73개 (20개 미달)",    "⚠️", "93+ 목표 미달성"),
    ("④무료범위추천",   "업무역량 향상 추천 문서",                        "cc4892f 내 예상 — 브랜치 미push로 불명",          "⛔", "의뢰 외 push 시도로 접근 불가"),
    ("⑤CLAUDE.md보강", "글로벌 CLAUDE.md 보강안",                       "e067e9d 내 예상 — 브랜치 미push로 불명",          "⛔", "idle timeout 영향 가능"),
    ("⑥VSCode통합",    "tasks.json + 슬래시명령 + 동기화 코드",          "tasks.json 10개 + commands 46개 확인",           "✅", "일부 UNREACH"),
    ("의뢰외 실행",     "git push 시도 (의뢰범위 초과)",                  "git push origin claude/create-korean-docs-ZzhQG", "🔴", "사용자 승인 없음 — 실패"),
    ("의뢰외 실행",     "GitHub MCP 자동호출 (쓰기성)",                   "MCP write tool 호출",                           "🔴", "사용자 승인 없음"),
    ("의뢰외 실행",     "API Stream idle timeout",                       "단일 대용량 배치 호출로 timeout 발생",            "🔴", "배치 분할 미적용"),
]
for r, row in enumerate(mapping, 2):
    for c, v in enumerate(row, 1):
        bg = LGRAY if r % 2 == 0 else WHITE
        fg = RED if v in ("🔴", "⛔") else GREEN if v == "✅" else "1E293B"
        data_cell(ws1, r, c, v, bg=bg, fg=fg, align="center" if c == 4 else "left")

ws1.column_dimensions["A"].width = 16
ws1.column_dimensions["B"].width = 28
ws1.column_dimensions["C"].width = 42
ws1.column_dimensions["D"].width = 10
ws1.column_dimensions["E"].width = 30

# ── Sheet2: 4축 매트릭스 ──
ws2 = wb.create_sheet("4축매트릭스")
ws2.sheet_view.showGridLines = False
ws2.row_dimensions[1].height = 32
hdrs2 = ["#", "의뢰범위", "파일경로", "①구현", "②실현", "③구동", "④검증", "비고"]
for i, h in enumerate(hdrs2, 1):
    hdr_cell(ws2, 1, i, h)

for r, f in enumerate(FILES, 2):
    ws2.row_dimensions[r].height = 28
    bg = LGRAY if r % 2 == 0 else WHITE
    data_cell(ws2, r, 1, f["id"],    bg=bg, align="center")
    data_cell(ws2, r, 2, f["scope"], bg=bg)
    data_cell(ws2, r, 3, f["path"][:60],  bg=bg)
    for col, key in [(4, "impl"), (5, "place"), (6, "run"), (7, "valid")]:
        v = f[key]
        fg = GREEN if v in ("PASS","NA") else RED if v == "FAIL" else AMBER
        data_cell(ws2, r, col, v, bg=bg, fg=fg, bold=True, align="center")
    data_cell(ws2, r, 8, f["note"], bg=bg)

ws2.column_dimensions["A"].width = 5
ws2.column_dimensions["B"].width = 16
ws2.column_dimensions["C"].width = 45
for col in "DEFG": ws2.column_dimensions[col].width = 10
ws2.column_dimensions["H"].width = 50

# ── Sheet3: 보완점 ──
ws3 = wb.create_sheet("보완점")
ws3.sheet_view.showGridLines = False
ws3.row_dimensions[1].height = 30
hdrs3 = ["항목", "구분", "원인", "영향도", "수정안", "예상시간"]
for i, h in enumerate(hdrs3, 1):
    hdr_cell(ws3, 1, i, h)

improvements = [
    ("브랜치 미push",        "FAIL/UNREACH", "git push 의뢰 외 자동실행 실패",            "상", "노트북에서 git push origin claude/create-korean-docs-ZzhQG", "2분"),
    ("CLAUDE.md 안전장치",   "설계 부재",    "push·MCP·timeout 방지 룰 미정의",           "상", "RULE-G1/G2/G3 3조 추가 (글로벌+워크스페이스)",              "10분"),
    ("에이전트 73→93+",      "PARTIAL",      "20개 미달 — 의뢰 목표 달성 불명",           "중", "추가 20개 agent proposal 신설 + active 전환 로드맵",         "2시간"),
    ("AI컨설팅시스템 docx",  "UNREACH",      "브랜치 미push or idle timeout으로 미생성",  "중", "브랜치 push 후 재확인, 미존재 시 OFFICE-B 신규 생성",        "1시간"),
    ("corp-financial 스킬",  "FAIL",         "전용 SKILL.md 파일 미신설",                 "하", "skills/법인재무분석/SKILL.md 신설",                           "20분"),
    ("idle timeout 방지",    "리스크",       "단일 대용량 호출 패턴 유지",                "중", "2,000토큰 단위 배치 분할 패턴 CLAUDE.md에 명시",             "30분"),
]
for r, row in enumerate(improvements, 2):
    bg = LGRAY if r % 2 == 0 else WHITE
    for c, v in enumerate(row, 1):
        fg = RED if v in ("상", "FAIL/UNREACH") else AMBER if v == "중" else "1E293B"
        data_cell(ws3, r, c, v, bg=bg, fg=fg, align="center" if c in (4, 6) else "left")

for col, w in zip("ABCDEF", [18, 14, 36, 8, 44, 10]):
    ws3.column_dimensions[col].width = w

# ── Sheet4: 체크리스트 ──
ws4 = wb.create_sheet("체크리스트")
ws4.sheet_view.showGridLines = False
ws4.row_dimensions[1].height = 28
hdrs4 = ["#", "우선순위", "조치 항목", "기한", "상태"]
for i, h in enumerate(hdrs4, 1):
    hdr_cell(ws4, 1, i, h)

checks = [
    (1, "긴급", "노트북에서 브랜치 push",                    "오늘",   "미완"),
    (2, "긴급", "CLAUDE.md 안전장치 3조 추가",               "오늘",   "미완"),
    (3, "높음", "브랜치 pull + 25파일 재검증",               "D+3",    "미완"),
    (4, "높음", "에이전트 93+ 달성 로드맵 수립",             "D+3",    "미완"),
    (5, "중간", "AI컨설팅시스템_종합역량문서_v1.docx 복원",  "D+7",    "미완"),
    (6, "중간", "corp-financial-analysis SKILL.md 신설",     "D+7",    "미완"),
    (7, "낮음", "VS Code task /컨설팅시작 원클릭 추가",       "D+14",   "미완"),
    (8, "낮음", "슬래시명령 /진단요약 신설",                  "D+14",   "미완"),
]
for r, row in enumerate(checks, 2):
    bg = LGRAY if r % 2 == 0 else WHITE
    fg_pri = RED if row[1]=="긴급" else AMBER if row[1]=="높음" else "1E293B"
    for c, v in enumerate(row, 1):
        fg = fg_pri if c == 2 else "1E293B"
        data_cell(ws4, r, c, v, bg=bg, fg=fg, align="center" if c in (1, 4, 5) else "left")

ws4.column_dimensions["A"].width = 5
ws4.column_dimensions["B"].width = 10
ws4.column_dimensions["C"].width = 48
ws4.column_dimensions["D"].width = 10
ws4.column_dimensions["E"].width = 10

wb.save(OUT_DIR / "검증결과_v1.xlsx")
print("✅ 산출물④: 검증결과_v1.xlsx 생성")


# ─────────────────────────────────────────────
# 산출물②: 검증보고서_v1.docx
# ─────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="000080"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.name = "Malgun Gothic"
    return p

def add_para(doc, text, sz=10, color="1E293B", bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sz)
    run.font.name = "Malgun Gothic"
    run.font.bold = bold
    if color != "1E293B":
        run.font.color.rgb = RGBColor.from_string(color)
    return p

# 표제부
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("검증 보고서 v1")
run.font.size = Pt(22); run.font.bold = True
run.font.color.rgb = RGBColor.from_string("000080")
run.font.name = "Malgun Gothic"

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("OFFICE-B 동기화 검증 | 2026-05-04 | 여운식 전문위원")
run2.font.size = Pt(10); run2.font.name = "Malgun Gothic"
run2.font.color.rgb = RGBColor.from_string("64748B")
doc.add_paragraph()

# Section 1
add_heading(doc, "1. 개요 및 핵심 결론")
alert = doc.add_paragraph()
r_a = alert.add_run("⚠️ 브랜치 claude/create-korean-docs-ZzhQG = 원격 미push — OFFICE-B 접근 불가")
r_a.font.color.rgb = RGBColor.from_string("DC2626")
r_a.font.bold = True; r_a.font.size = Pt(10); r_a.font.name = "Malgun Gothic"

add_para(doc, "노트북(LAPTOP/sik37) Claude Code 세션이 의뢰 ①~⑥을 처리한 결과, 로컬 커밋 2건(cc4892f, e067e9d)이 생성되었으나 원격 push가 미완료되어 OFFICE-B에서 해당 파일에 접근할 수 없습니다.", sz=9)
add_para(doc, "또한 의뢰 범위('코드 생성까지')를 초과하여 git push 시도 및 GitHub MCP 쓰기성 도구 자동호출이 발생한 것이 확인되었습니다.", sz=9)

# Section 2: 매핑표
add_heading(doc, "2. 의뢰범위 ①~⑥ vs 실제실행 매핑표")
tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = "Table Grid"
hdrs = ["의뢰", "내용", "실제 결과", "의뢰범위 내"]
for i, h in enumerate(hdrs):
    cell = tbl2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].runs[0].font.name = "Malgun Gothic"
    set_cell_bg(cell, "000080")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")

map_data = [
    ("①코워크모드스킬", "30개 스킬 플러그인",         "skills/ 30개 SKILL.md ✅",      "✅"),
    ("②플러그인전수",   "cowork 구성",                 "PLUGIN_MAPPING.md + README ✅",  "✅"),
    ("③에이전트93+",    "93개 이상 전수",              "73개 (active33+proposals40) ⚠️", "⚠️"),
    ("④무료범위추천",   "업무역량 향상 문서",           "브랜치 미push — 접근 불가 ⛔",   "🔴"),
    ("⑤CLAUDE.md보강", "글로벌 CLAUDE.md 업데이트",   "브랜치 미push — 불명 ⛔",        "🔴"),
    ("⑥VSCode통합",    "tasks.json + 슬래시명령",      "tasks 10개 + commands 46개 ✅",  "✅"),
]
for row_data in map_data:
    row = tbl2.add_row()
    for i, v in enumerate(row_data):
        row.cells[i].text = v
        run = row.cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9); run.font.name = "Malgun Gothic"
        if "🔴" in v or "⛔" in v:
            run.font.color.rgb = RGBColor.from_string("DC2626")
        elif "⚠️" in v:
            run.font.color.rgb = RGBColor.from_string("F59E0B")
        elif "✅" in v:
            run.font.color.rgb = RGBColor.from_string("047857")

doc.add_paragraph()

# Section 3: 4축 매트릭스
add_heading(doc, "3. 4축 검증 매트릭스 (25파일)")
tbl3 = doc.add_table(rows=1, cols=7)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.style = "Table Grid"
h3 = ["#", "범위", "파일", "①구현", "②실현", "③구동", "④검증"]
for i, h in enumerate(h3):
    c = tbl3.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].font.bold = True
    c.paragraphs[0].runs[0].font.size = Pt(8)
    c.paragraphs[0].runs[0].font.name = "Malgun Gothic"
    set_cell_bg(c, "000080")
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")

STATUS_COL = {"PASS": "047857", "FAIL": "DC2626", "NA": "64748B", "UNREACH": "F59E0B"}
for f in FILES:
    row = tbl3.add_row()
    vals = [str(f["id"]), f["scope"][:8], f["path"].split("/")[-1][:22],
            f["impl"], f["place"], f["run"], f["valid"]]
    for i, v in enumerate(vals):
        cell = row.cells[i]
        cell.text = v
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(7.5); run.font.name = "Malgun Gothic"
        if v in STATUS_COL:
            run.font.color.rgb = RGBColor.from_string(STATUS_COL[v])
            run.font.bold = True

doc.add_paragraph()

# Section 4: 의뢰 외 실행
add_heading(doc, "4. 🔴 의뢰 외 자동실행 (경계 위반)")
for v in VIOLATIONS:
    p = doc.add_paragraph(style="List Bullet")
    r_title = p.add_run(f"[{v['no']}] {v['action']}: ")
    r_title.font.color.rgb = RGBColor.from_string("DC2626")
    r_title.font.bold = True; r_title.font.size = Pt(9); r_title.font.name = "Malgun Gothic"
    r_detail = p.add_run(f"{v['scope']} → {v['impact']}")
    r_detail.font.size = Pt(9); r_detail.font.name = "Malgun Gothic"

doc.add_paragraph()

# Section 5: 보완점
add_heading(doc, "5. 보완점 및 안전장치 설계")
add_para(doc, "5-1. CLAUDE.md 안전장치 3조 (즉시 추가 권고)", sz=10, bold=True, color="000080")
for rule in [
    "RULE-G1: 명시 없는 git push 금지 (검증 완료 + 사용자 승인 후에만)",
    "RULE-G2: MCP 쓰기성 도구 자동호출 시 사용자 승인 필수",
    "RULE-G3: API Stream idle timeout 시 배치 분할 재시도 (2,000토큰 단위, 최대 3회)",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(rule)
    r.font.size = Pt(9); r.font.name = "Malgun Gothic"
    r.font.color.rgb = RGBColor.from_string("DC2626")

doc.add_paragraph()
add_para(doc, "5-2. 무료 범위 추가 보강 후보 3건", sz=10, bold=True, color="000080")
for item in [
    "① VS Code task '/컨설팅시작' 원클릭 (consulting-agent 실행 자동화)",
    "② 슬래시명령 /진단요약 신설 (4자관점 1페이지 진단 요약 자동 생성)",
    "③ skills/법인재무분석/SKILL.md 신설 (코워크모드 직접 호출)",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(item)
    r.font.size = Pt(9); r.font.name = "Malgun Gothic"

doc.save(OUT_DIR / "검증보고서_v1.docx")
print("✅ 산출물②: 검증보고서_v1.docx 생성")

# ─────────────────────────────────────────────
# 최종 확인
# ─────────────────────────────────────────────
print("\n=== 5종 산출물 생성 완료 ===")
for fname in [f"검증_1페이지요약_{DATE_STR}.md", "검증보고서_v1.docx",
              "검증대시보드_v1.html", "검증결과_v1.xlsx", "보완_체크리스트_v1.md"]:
    fpath = OUT_DIR / fname
    exists = fpath.exists()
    size   = fpath.stat().st_size if exists else 0
    print(f"  {'✅' if exists else '❌'} {fname} ({size:,} bytes)")
