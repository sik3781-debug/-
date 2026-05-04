# -*- coding: utf-8 -*-
"""
gen_verify_v2.py — Phase 4 검증 산출물 5종 생성
실행: python outputs/gen_verify_v2.py
"""
import pathlib, openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(r"C:\Users\Jy\consulting-agent")
OUT  = ROOT / "outputs"

# ──────────────────────────────────────
# 31건 검증 결과 데이터
# ──────────────────────────────────────
RESULTS = [
    (1,"①스킬",".skills/corp-financial-analysis/SKILL.md","PASS","PASS","PASS","PASS",""),
    (2,"①스킬",".skills/pptx/SKILL.md","PASS","PASS","PASS","PASS",""),
    (3,"①스킬",".skills/docx/SKILL.md","PASS","PASS","PASS","PASS",""),
    (4,"①스킬",".skills/xlsx/SKILL.md","PASS","PASS","PASS","PASS",""),
    (5,"①스킬",".skills/pdf/SKILL.md","PASS","PASS","PASS","PASS",""),
    (6,"②cmd",".claude/commands/마스터컨설팅.md","PASS","PASS","PASS","PASS",""),
    (7,"②cmd",".claude/commands/비상장주식평가.md","PASS","PASS","PASS","PASS",""),
    (8,"②cmd",".claude/commands/가지급금해소시뮬.md","PASS","PASS","PASS","PASS",""),
    (9,"②cmd",".claude/commands/가업승계시뮬.md","PASS","PASS","PASS","PASS",""),
    (10,"②cmd",".claude/commands/차명주식해소.md","PASS","PASS","PASS","PASS",""),
    (11,"②cmd",".claude/commands/임원퇴직금한도.md","PASS","PASS","PASS","PASS",""),
    (12,"②cmd",".claude/commands/상속증여시뮬.md","PASS","PASS","PASS","PASS",""),
    (13,"②cmd",".claude/commands/특허패키지.md","PASS","PASS","PASS","PASS",""),
    (14,"②cmd",".claude/commands/금융기관진단.md","PASS","PASS","PASS","PASS",""),
    (15,"②cmd",".claude/commands/법인재무종합분석.md","PASS","PASS","PASS","PASS",""),
    (16,"③vscode",".vscode/tasks.json","PASS","PASS","PASS","PASS",""),
    (17,"③vscode",".vscode/launch.json","PASS","PASS","PASS","PASS",""),
    (18,"③vscode",".vscode/settings.json","PASS","PASS","PASS","PASS",""),
    (19,"③vscode",".vscode/extensions.json","PASS","PASS","PASS","PASS",""),
    (20,"③vscode","setup_vscode_integration.py","PASS","PASS","PASS","PASS",""),
    (21,"④sync","sync_all_devices.ps1","PASS","PASS","PASS","PASS","UTF-8 BOM 정상"),
    (22,"④sync","sync_all_devices.sh","PASS","PASS","NA","PASS","sh — 구동 NA(Windows)"),
    (23,"⑤doc","outputs/시스템_종합역량_문서_v1.md","PASS","PASS","PASS","PASS","742줄 5챕터"),
    (24,"⑤doc","outputs/AI컨설팅시스템_종합역량문서_v1.docx","PASS","PASS","PASS","NA","5챕터 구조 정상"),
    (25,"⑥claude","CLAUDE.md","PASS","PASS","PASS","PASS","다기기+VSCode 섹션 보강"),
    (26,"⑥claude","CLAUDE.md(RULE-G 3조)","PASS","PASS","NA","PASS","G1·G2·G3 3 hit"),
    (27,"⑦infra",".vscode/tasks.json (10tasks)","PASS","PASS","PASS","PASS","tasks.Count=10"),
    (28,"⑦infra",".skills/ (5종)","PASS","PASS","NA","PASS","SKILL.md 5건"),
    (29,"⑦infra",".claude/commands/ (10종)","PASS","PASS","NA","PASS","*.md 10건"),
    (30,"⑦infra","outputs/gen_docx_v1.py","PASS","PASS","PASS","PASS",""),
    (31,"⑦infra","outputs/gen_verification_outputs.py","PASS","PASS","PASS","PASS","Auto-Fix 1회"),
]

total = len(RESULTS)
pass_all = sum(1 for r in RESULTS if r[3]=="PASS" and r[4]=="PASS" and r[5] in ("PASS","NA") and r[6] in ("PASS","NA"))
fail_c   = sum(1 for r in RESULTS if "FAIL" in r[3:7])
pass_pct = round(pass_all/total*100)

# ───────────────────────────────────────
# ① 1페이지 요약 .md
# ───────────────────────────────────────
summary = f"""# 검증 1페이지 요약 v2 — OFFICE-B 2026-05-04

**검증 기준**: Phase 0~4 완료 후 31건 4축 매트릭스
**Auto-Fix**: 1회 적용 (#31 gen_verification_outputs.py 복사)

## 4축 매트릭스 집계

| 결과 | 건수 | 비율 |
|---|---|---|
| ✅ ALL PASS | {pass_all} | {pass_pct}% |
| ❌ FAIL | {fail_c} | {100-pass_pct}% |
| **합계** | **{total}** | **100%** |

## 카테고리별 결과

| 카테고리 | 파일 수 | 결과 |
|---|---|---|
| ① .skills/ | 5 | ✅ ALL PASS |
| ② .claude/commands/ | 10 | ✅ ALL PASS |
| ③ .vscode/ + setup | 5 | ✅ ALL PASS |
| ④ sync 스크립트 | 2 | ✅ PASS (sh 구동 NA) |
| ⑤ 종합문서 md+docx | 2 | ✅ PASS |
| ⑥ CLAUDE.md 보강 | 2 | ✅ PASS |
| ⑦ 인프라·생성스크립트 | 5 | ✅ ALL PASS (Auto-Fix 1) |

## RULE-G 자가검증

| 질문 | 답 |
|---|---|
| 사용자 명시 없는 push 시도 | ✅ 0회 |
| MCP 쓰기성 도구 무승인 호출 | ✅ 0회 |
| 단일 호출 5파일·2,000토큰 초과 | ✅ 0회 (배치 분할 준수) |
| VS Code 통합 Claude Code 환경 | ✅ DESKTOP-FJUATON/Jy |

## 잔재 정리 (Phase 6) 승인 요청 대기
- 임시 파일 A~F 카테고리 식별 완료
- 사용자 응답 대기: "전부 삭제" / "선택적 삭제" / "보류"

*생성: 2026-05-04 OFFICE-B Phase 4 Batch 7/7*
"""
(OUT / "검증_1페이지요약_20260504_v2.md").write_text(summary, encoding="utf-8-sig")
print("✅ 산출물①: 검증_1페이지요약_20260504_v2.md")

# ───────────────────────────────────────
# ② 검증보고서 v2 .docx
# ───────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8); s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("검증 보고서 v2"); r.font.size = Pt(20); r.font.bold = True
r.font.color.rgb = RGBColor.from_string("000080"); r.font.name = "Malgun Gothic"
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(f"OFFICE-B | 2026-05-04 | 31건 4축 매트릭스 | {pass_all}/31 PASS")
r2.font.size = Pt(9); r2.font.name = "Malgun Gothic"
r2.font.color.rgb = RGBColor.from_string("64748B")
doc.add_paragraph()

ph = doc.add_heading("4축 검증 매트릭스 (31건)", level=1)
for r3 in ph.runs: r3.font.color.rgb = RGBColor.from_string("000080"); r3.font.name = "Malgun Gothic"

tbl = doc.add_table(rows=1, cols=7); tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["#","범위","파일","①구현","②실현","③구동","④검증"]):
    c = tbl.rows[0].cells[i]; c.text = h
    r = c.paragraphs[0].runs[0]; r.font.bold=True; r.font.size=Pt(8); r.font.name="Malgun Gothic"
    r.font.color.rgb = RGBColor.from_string("FFFFFF"); set_bg(c, "000080")
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

STATUS_COL = {"PASS":"047857","FAIL":"DC2626","NA":"64748B"}
for res in RESULTS:
    row = tbl.add_row()
    vals = [str(res[0]), res[1], res[2][-25:], res[3], res[4], res[5], res[6]]
    for i, v in enumerate(vals):
        cell = row.cells[i]; cell.text = v
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(7.5); run.font.name = "Malgun Gothic"
        if v in STATUS_COL:
            run.font.color.rgb = RGBColor.from_string(STATUS_COL[v]); run.font.bold = True

doc.add_paragraph()
ph2 = doc.add_heading("Auto-Fix 적용 내역", level=1)
for r3 in ph2.runs: r3.font.color.rgb = RGBColor.from_string("000080"); r3.font.name = "Malgun Gothic"
p3 = doc.add_paragraph()
r3 = p3.add_run("#31 gen_verification_outputs.py — junggi-workspace/outputs/ → consulting-agent/outputs/ 복사 (1회, PASS)")
r3.font.size = Pt(9); r3.font.name = "Malgun Gothic"

doc.save(OUT / "검증보고서_v2.docx")
print("✅ 산출물②: 검증보고서_v2.docx")

# ───────────────────────────────────────
# ③ 검증 대시보드 v2 .html
# ───────────────────────────────────────
file_rows = ""
for res in RESULTS:
    cls = lambda v: "pass" if v in ("PASS","NA") else "fail"
    note = res[7][:40] if res[7] else ""
    file_rows += f"""<tr>
<td style="text-align:center">{res[0]}</td>
<td style="font-size:11px">{res[1]}</td>
<td style="font-size:10px">{res[2].split('/')[-1][:24]}</td>
<td class="{cls(res[3])}">{res[3]}</td>
<td class="{cls(res[4])}">{res[4]}</td>
<td class="{cls(res[5])}">{res[5]}</td>
<td class="{cls(res[6])}">{res[6]}</td>
<td style="font-size:9px;color:#64748b">{note}</td>
</tr>"""

html = f"""<!DOCTYPE html><html lang="ko"><head><meta charset="UTF-8">
<title>검증 대시보드 v2 — OFFICE-B 2026-05-04</title>
<style>
body{{font-family:'Malgun Gothic',sans-serif;margin:20px;background:#f8fafc;color:#1e293b}}
h1{{color:#000080;font-size:20px}}h2{{color:#1e40af;font-size:14px;border-bottom:2px solid #d1d9e6;padding-bottom:4px}}
.kpi-row{{display:flex;gap:12px;margin:12px 0}}
.kpi{{background:#fff;border-radius:8px;padding:12px 20px;border-left:5px solid #000080;box-shadow:0 1px 4px rgba(0,0,0,.08)}}
.kpi.fail{{border-left-color:#dc2626}}.kpi-num{{font-size:28px;font-weight:700}}.kpi-label{{font-size:11px;color:#64748b}}
svg.donut{{width:100px;height:100px}}.donut-wrap{{display:flex;align-items:center;gap:16px;background:#fff;border-radius:8px;padding:12px;box-shadow:0 1px 4px rgba(0,0,0,.08);width:fit-content}}
.legend{{font-size:12px;line-height:2}}.dot{{display:inline-block;width:10px;height:10px;border-radius:50%;margin-right:4px}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:8px;overflow:hidden;box-shadow:0 1px 4px rgba(0,0,0,.08);font-size:11px}}
th{{background:#000080;color:#fff;padding:7px;text-align:center}}td{{padding:5px 7px;border-bottom:1px solid #f1f5f9}}
tr:hover td{{background:#f0f4ff}}td.pass{{color:#047857;font-weight:700;text-align:center}}td.fail{{color:#dc2626;font-weight:700;text-align:center}}
.badge{{background:#047857;color:#fff;border-radius:4px;padding:2px 6px;font-size:11px}}
</style></head><body>
<h1>검증 대시보드 v2 — OFFICE-B / 2026-05-04</h1>
<p>31건 재현 + 4축 매트릭스 + Auto-Fix 1회 | RULE-G 자가검증 3축 ALL ✅</p>
<div class="kpi-row">
<div class="kpi"><div class="kpi-num" style="color:#047857">{pass_all}</div><div class="kpi-label">✅ ALL PASS</div></div>
<div class="kpi fail"><div class="kpi-num" style="color:#dc2626">{fail_c}</div><div class="kpi-label">❌ FAIL</div></div>
<div class="kpi"><div class="kpi-num">{total}</div><div class="kpi-label">전체 파일</div></div>
<div class="kpi"><div class="kpi-num" style="color:#000080">{pass_pct}%</div><div class="kpi-label">PASS 비율</div></div>
</div>
<h2>PASS 분포 도넛</h2>
<div class="donut-wrap">
<svg class="donut" viewBox="0 0 42 42">
<circle cx="21" cy="21" r="15.91" fill="transparent" stroke="#047857" stroke-width="6"
  stroke-dasharray="{pass_pct} {100-pass_pct}" stroke-dashoffset="25"/>
<circle cx="21" cy="21" r="15.91" fill="transparent" stroke="#dc2626" stroke-width="6"
  stroke-dasharray="{100-pass_pct} {pass_pct}" stroke-dashoffset="{125-pass_pct}"/>
<text x="21" y="21" text-anchor="middle" dy=".35em" font-size="7" fill="#000080" font-weight="bold">{pass_pct}%</text>
</svg>
<div class="legend">
<div><span class="dot" style="background:#047857"></span>PASS {pass_all}건 ({pass_pct}%)</div>
<div><span class="dot" style="background:#dc2626"></span>FAIL {fail_c}건 ({100-pass_pct}%)</div>
</div></div>
<h2>4축 매트릭스 (31파일)</h2>
<table><thead><tr><th>#</th><th>범위</th><th>파일</th><th>①구현</th><th>②실현</th><th>③구동</th><th>④검증</th><th>비고</th></tr></thead>
<tbody>{file_rows}</tbody></table>
<h2>RULE-G 자가검증</h2>
<table><thead><tr><th>질문</th><th>답변</th></tr></thead><tbody>
<tr><td>사용자 명시 없는 push 시도</td><td class="pass">✅ 0회</td></tr>
<tr><td>MCP 쓰기성 도구 무승인 호출</td><td class="pass">✅ 0회</td></tr>
<tr><td>단일 호출 5파일·2,000토큰 초과</td><td class="pass">✅ 0회 (배치7개 분할)</td></tr>
<tr><td>VS Code 통합 Claude Code 환경</td><td class="pass">✅ DESKTOP-FJUATON/Jy</td></tr>
</tbody></table>
<footer style="margin-top:24px;font-size:10px;color:#64748b">생성: 2026-05-04 OFFICE-B | Claude Sonnet 4.6</footer>
</body></html>"""

(OUT / "검증대시보드_v2.html").write_text(html, encoding="utf-8-sig")
print("✅ 산출물③: 검증대시보드_v2.html")

# ───────────────────────────────────────
# ④ 검증결과 v2 .xlsx
# ───────────────────────────────────────
wb = openpyxl.Workbook()
wb.remove(wb.active)
thin = Side(style="thin", color="D1D9E6")
bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)

def hdr(ws, r, c, v, bg="000080", fg="FFFFFF"):
    cell = ws.cell(r, c, v)
    cell.fill = PatternFill("solid", fgColor=bg)
    cell.font = Font(color=fg, bold=True, size=9, name="Calibri")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    cell.border = bdr

def dat(ws, r, c, v, bg="FFFFFF", fg="1E293B", align="left"):
    cell = ws.cell(r, c, v)
    cell.fill = PatternFill("solid", fgColor=bg)
    cell.font = Font(color=fg, size=9, name="Calibri")
    cell.alignment = Alignment(horizontal=align, vertical="center")
    cell.border = bdr

SC = {"PASS":"047857","FAIL":"DC2626","NA":"64748B"}

# Sheet1: 매트릭스
ws1 = wb.create_sheet("4축매트릭스")
ws1.sheet_view.showGridLines = False
for i, h in enumerate(["#","범위","파일경로","①구현","②실현","③구동","④검증","비고"], 1):
    hdr(ws1, 1, i, h)
for ri, res in enumerate(RESULTS, 2):
    bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
    dat(ws1, ri, 1, res[0], bg, align="center")
    dat(ws1, ri, 2, res[1], bg)
    dat(ws1, ri, 3, res[2], bg)
    for ci, v in enumerate([res[3],res[4],res[5],res[6]], 4):
        dat(ws1, ri, ci, v, bg, SC.get(v,"1E293B"), "center")
    dat(ws1, ri, 8, res[7], bg)

for col, w in zip("ABCDEFGH", [5,10,40,10,10,10,10,30]):
    ws1.column_dimensions[col].width = w

# Sheet2: 요약
ws2 = wb.create_sheet("요약")
ws2.sheet_view.showGridLines = False
hdr(ws2, 1, 1, "항목"); hdr(ws2, 1, 2, "값")
summary_data = [
    ("전체 파일", total), ("ALL PASS", pass_all), ("FAIL", fail_c),
    ("PASS 비율", f"{pass_pct}%"), ("Auto-Fix 횟수", 1),
    ("RULE-G push 위반", 0), ("RULE-G MCP 위반", 0), ("RULE-G 배치 위반", 0),
    ("Batch 커밋 수", 7), ("총 커밋 (RULE-G포함)", 8),
]
for ri, (k, v) in enumerate(summary_data, 2):
    bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
    dat(ws2, ri, 1, k, bg); dat(ws2, ri, 2, str(v), bg, align="center")
ws2.column_dimensions["A"].width = 25; ws2.column_dimensions["B"].width = 15

# Sheet3: 체크리스트
ws3 = wb.create_sheet("체크리스트")
ws3.sheet_view.showGridLines = False
for i, h in enumerate(["#","우선순위","항목","기한","상태"], 1):
    hdr(ws3, 1, i, h)
checks = [
    (1,"긴급","push 승인 요청 (8 commit → origin/main)","오늘","대기"),
    (2,"높음","노트북·HOME-A pull + 재검증","D+1","대기"),
    (3,"높음","에이전트 93+ 달성 로드맵","D+7","미완"),
    (4,"중간","/진단요약 슬래시명령 신설","D+7","미완"),
    (5,"중간","법인재무분석 SKILL.md 신설","D+7","미완"),
    (6,"낮음","pre_session_check.ps1 신설","D+14","미완"),
]
for ri, row in enumerate(checks, 2):
    bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
    fg_p = "DC2626" if row[1]=="긴급" else "F59E0B" if row[1]=="높음" else "1E293B"
    for ci, v in enumerate(row, 1):
        dat(ws3, ri, ci, v, bg, fg_p if ci == 2 else "1E293B", "center" if ci in (1,4,5) else "left")
for col, w in zip("ABCDE", [5,10,42,10,10]):
    ws3.column_dimensions[col].width = w

wb.save(OUT / "검증결과_v2.xlsx")
print("✅ 산출물④: 검증결과_v2.xlsx")

# ───────────────────────────────────────
# ⑤ 보완 체크리스트 v2
# ───────────────────────────────────────
checklist = f"""# 보완 체크리스트 v2 — 자동수정 적용 결과 (2026-05-04)

## Auto-Fix 적용 결과

| # | 항목 | Auto-Fix | 결과 |
|---|---|---|---|
| 31 | gen_verification_outputs.py 위치 오류 | junggi-workspace → consulting-agent 복사 | ✅ PASS |

**Auto-Fix 총계**: 1건 / 1회 시도 / 100% 성공률

---

## A. 즉시 조치 (오늘)

- [ ] **push 승인 요청** — 8개 commit(RULE-G + Batch 1~7) → origin/main
  - 영향도: 상 | 예상시간: 1분 (승인 후 즉시)
  - 명령: `git push origin main`

- [ ] **노트북·HOME-A pull** — 이 기기 push 완료 후 양 기기에서 pull
  - 영향도: 상 | 예상시간: 2분/기기

## B. 단기 조치 (D+7)

- [ ] **/진단요약 슬래시명령 신설**
  - `.claude/commands/진단요약.md`
  - 영향도: 중 | 예상시간: 15분

- [ ] **법인재무분석 SKILL.md 신설**
  - `.skills/법인재무분석/SKILL.md`
  - 영향도: 중 | 예상시간: 20분

- [ ] **에이전트 93+ 달성 로드맵 착수**
  - 현재 73개 → +20개 (D+45 완료 목표)
  - 영향도: 중 | 예상시간: 2시간/배치

## C. 중기 조치 (D+14)

- [ ] **pre_session_check.ps1 신설**
  - `bootstrap/pre_session_check.ps1`
  - 영향도: 낮음 | 예상시간: 30분

- [ ] **DART API 발급 후 통합 활성화**
  - DART_API_KEY, LAW_API_ID, ECOS_API_KEY, PUBLIC_DATA_API_KEY
  - 영향도: 중 | 예상시간: 발급 후 30분

---

## RULE-G 자가검증 최종

| 질문 | 답변 |
|---|---|
| 사용자 명시 없는 push 시도 0회? | ✅ Y |
| MCP 쓰기성 도구 무승인 호출 0회? | ✅ Y |
| 단일 호출 5파일·2,000토큰 초과 0회? | ✅ Y (7배치 분할) |
| VS Code 통합 Claude Code 환경? | ✅ Y (DESKTOP-FJUATON/Jy) |

**4개 모두 Y → 정상 종료 조건 충족**

---

## 잔재 정리 Phase 6 승인 요청

삭제 후보:
- A: 임시 산출물 (_pre_snapshot.txt 등 5건)
- B: backup-before-redo 태그 1건
- E: __pycache__·*.pyc 빌드 잔재

다음 중 하나로 응답해 주십시오:
- **"전부 삭제 진행"**
- **"선택적 삭제 (A만)"**
- **"잔재 정리 보류"**

*생성: 2026-05-04 OFFICE-B Phase 4 Batch 7/7*
"""

(OUT / "보완_체크리스트_v2_적용결과.md").write_text(checklist, encoding="utf-8-sig")
print("✅ 산출물⑤: 보완_체크리스트_v2_적용결과.md")

print(f"\n=== 5종 산출물 생성 완료 ({pass_all}/{total} PASS) ===")
for fn in ["검증_1페이지요약_20260504_v2.md","검증보고서_v2.docx",
           "검증대시보드_v2.html","검증결과_v2.xlsx","보완_체크리스트_v2_적용결과.md"]:
    p = OUT / fn
    print(f"  ✅ {fn} ({p.stat().st_size:,} bytes)")
